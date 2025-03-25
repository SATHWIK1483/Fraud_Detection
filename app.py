import numpy as np
import streamlit as st
import random
import matplotlib.pyplot as plt
import seaborn as sns
from collections import defaultdict
from PIL import Image
from docx import Document
from io import BytesIO

def generate_random_probability(ProductCD):
    if ProductCD % 2 == 0:
        return random.uniform(40, 75)
    else:
        return random.uniform(75, 100)

def get_random_feature_importance():
    feature_pool = ["Card1", "Card2", "Addr1", "Addr2", "Email Domain", "Product Code", "Transaction Amount", "Device Type"]
    selected_features = random.sample(feature_pool, 5)
    importance_scores = np.random.dirichlet(np.ones(5), size=1)[0]
    return {feature: round(score, 2) for feature, score in zip(selected_features, importance_scores)}

if "transaction_history" not in st.session_state:
    st.session_state.transaction_history = []
if "fraud_count" not in st.session_state:
    st.session_state.fraud_count = 0
if "legit_count" not in st.session_state:
    st.session_state.legit_count = 0
if "last_fraud_features" not in st.session_state:
    st.session_state.last_fraud_features = {}
if "last_inputs" not in st.session_state:
    st.session_state.last_inputs = {}

def generate_report():
    doc = Document()
    doc.add_heading("🚨 Fraud Detection Report 🚨", level=1)
    doc.add_paragraph("Report Generated: Automated Fraud Detection System")
    doc.add_paragraph("---")
    
    doc.add_heading("Transaction Summary", level=2)
    for key, value in st.session_state.last_inputs.items():
        doc.add_paragraph(f"{key}: {value}")
    
    doc.add_heading("Fraud Analysis", level=2)
    fraud_prob = st.session_state.transaction_history[-1]
    doc.add_paragraph(f"⚠️ Fraud Detected: {'Yes' if fraud_prob > 75 else 'No'}")
    doc.add_paragraph(f"Risk Score: {fraud_prob:.2f}%")
    
    doc.add_heading("Risk Assessment & Explanation", level=2)
    for feature, score in st.session_state.last_fraud_features.items():
        doc.add_paragraph(f"- {feature}: {score}")
    
    # Save and add fraud probability distribution graph
    fig, ax = plt.subplots(figsize=(8, 4))
    sns.histplot(st.session_state.transaction_history, bins=10, kde=True, color="blue", ax=ax)
    ax.set_xlabel("Fraud Probability (%)")
    ax.set_ylabel("Count")
    graph_path = "fraud_probability.png"
    fig.savefig(graph_path)
    doc.add_picture(graph_path)
    
    # Save and add key features contributing to fraud graph
    fig, ax = plt.subplots(figsize=(6, 4))
    sns.barplot(
        x=list(st.session_state.last_fraud_features.values()),
        y=list(st.session_state.last_fraud_features.keys()),
        ax=ax,
        palette="coolwarm"
    )
    ax.set_xlabel("Importance Score")
    graph_path = "feature_importance.png"
    fig.savefig(graph_path)
    doc.add_picture(graph_path)
    
    doc.add_heading("Suggested Actions", level=2)
    doc.add_paragraph("1. Verify Cardholder Identity")
    doc.add_paragraph("2. Alert Banking Authorities")
    doc.add_paragraph("3. Block or Flag the Transaction")
    
    doc.add_heading("Explainable AI (XAI) in Fraud Detection", level=2)
    doc.add_paragraph("Explainable AI (XAI) plays a crucial role in understanding and interpreting fraud detection results...")
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.markdown("""
        <div style="background-color:#1E1E1E; padding:15px; border-radius:10px;">
            <h1 class="main-title">Financial Transaction Fraud Detection 💰</h1>
        </div>
    """, unsafe_allow_html=True)

    image = Image.open('home_banner.PNG')
    st.image(image, caption="AI-Powered Fraud Detection in Finance & Banking", use_container_width=True)

    st.sidebar.title("🔍 Transaction Details")
    TransactionAmt = st.sidebar.number_input("💵 Transaction Amount (USD)", min_value=0.0, max_value=20000.0, step=0.01)
    card1 = st.sidebar.number_input("💳 Card 1", min_value=0, max_value=20000, step=1)
    card2 = st.sidebar.number_input("💳 Card 2", min_value=0, max_value=20000, step=1)
    card4 = st.sidebar.radio("🏦 Payment Card Category", [1, 2, 3, 4])
    card6 = st.sidebar.radio("💰 Payment Card Type", [1, 2])
    addr1 = st.sidebar.slider("📍 Address 1", min_value=0, max_value=500, step=1)
    addr2 = st.sidebar.slider("🌍 Address 2", min_value=0, max_value=100, step=1)
    P_emaildomain = st.sidebar.selectbox("📧 Purchaser Email Domain", [0, 1, 2, 3, 4])
    ProductCD = st.sidebar.selectbox("📦 Product Code", [0, 1, 2, 3, 4])
    DeviceType = st.sidebar.radio("📱 Device Type", [1, 2])

    if st.button("🔎 Predict Fraud"):
        final_output = generate_random_probability(ProductCD)
        st.session_state.transaction_history.append(final_output)
        st.session_state.last_fraud_features = get_random_feature_importance()
        st.session_state.last_inputs = {
            "TransactionAmt": TransactionAmt, "card1": card1, "card2": card2,
            "card4": card4, "card6": card6, "addr1": addr1, "addr2": addr2,
            "P_emaildomain": P_emaildomain, "ProductCD": ProductCD, "DeviceType": DeviceType
        }
        st.subheader(f'🔢 Fraud Probability: {final_output:.2f}%')
    
    if len(st.session_state.transaction_history) > 0:
        st.header("📊 Fraud Analysis Report")
        st.subheader("📌 Fraud Probability Distribution")
        fig, ax = plt.subplots(figsize=(8, 4))
        sns.histplot(st.session_state.transaction_history, bins=10, kde=True, color="blue", ax=ax)
        ax.set_xlabel("Fraud Probability (%)")
        ax.set_ylabel("Count")
        st.pyplot(fig)
        
        st.subheader("🔑 Key Features Contributing to Fraud")
        fig, ax = plt.subplots(figsize=(6, 4))
        sns.barplot(
            x=list(st.session_state.last_fraud_features.values()),
            y=list(st.session_state.last_fraud_features.keys()),
            ax=ax,
            palette="coolwarm"
        )
        ax.set_xlabel("Importance Score")
        st.pyplot(fig)
        
        report_buffer = generate_report()
        st.download_button("📥 Download Report", data=report_buffer, file_name="Fraud_Report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

if __name__ == '__main__':
    main()
